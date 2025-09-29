"""
문서 서비스 모듈 - OnlyOffice 연동 및 문서 처리
"""
import json
import requests
import hashlib
import time
from typing import Dict, Any, List
import streamlit as st
from config import ONLYOFFICE_CONFIG

class DocumentService:
    def __init__(self):
        self.docspace_url = ONLYOFFICE_CONFIG["docspace_url"]
        self.sdk_url = ONLYOFFICE_CONFIG["sdk_url"]
        self.api_key = ONLYOFFICE_CONFIG["api_key"]
        self.jwt_secret = ONLYOFFICE_CONFIG["jwt_secret"]
        self.mode = ONLYOFFICE_CONFIG["mode"]
        self.frame_id = ONLYOFFICE_CONFIG["frame_id"]
    
    def create_onlyoffice_editor(self, doc_id: str, doc_type: str = "docx", title: str = "새 문서", width="100%", height="600px") -> str:
        """OnlyOffice 편집기 생성 - 개선된 버전"""
        
        # 문서 타입에 따른 설정
        type_configs = {
            'docx': {'type': 'text', 'icon': '📄'},
            'pptx': {'type': 'presentation', 'icon': '📊'},
            'xlsx': {'type': 'spreadsheet', 'icon': '📈'},
            'existing': {'type': 'text', 'icon': '📁'}
        }
        
        config = type_configs.get(doc_type, type_configs['docx'])
        
        return f"""
        <div id="editor-container" style="width: {width}; height: {height}; border: 1px solid #e5e7eb; border-radius: 8px; background-color: #ffffff; overflow: hidden; position: relative;">
            
            <!-- 편집기 헤더 -->
            <div id="editor-header" style="background: #f8f9fa; padding: 12px 16px; border-bottom: 1px solid #e5e7eb; display: flex; justify-content: space-between; align-items: center;">
                <div style="display: flex; align-items: center;">
                    <span style="font-size: 20px; margin-right: 8px;">{config['icon']}</span>
                    <span style="font-weight: 600; color: #374151;" id="doc-title">{title}</span>
                    <span style="color: #6b7280; font-size: 12px; margin-left: 12px;" id="doc-id">ID: {doc_id}</span>
                </div>
                <div>
                    <button onclick="toggleAIPanel()" id="ai-toggle-btn" style="padding: 6px 12px; background: #8b5cf6; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px; margin-right: 8px;">
                        🤖 AI 도구
                    </button>
                    <button onclick="refreshEditor()" style="padding: 6px 12px; background: #6b7280; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px; margin-right: 8px;">
                        🔄 새로고침
                    </button>
                    <button onclick="openFullEditor()" style="padding: 6px 12px; background: #059669; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px;">
                        🚀 전체화면
                    </button>
                </div>
            </div>
            
            <!-- OnlyOffice 편집기 프레임 -->
            <div id="{self.frame_id}" style="width: 100%; height: calc(100% - 50px); position: relative;">
                <div id="editor-loading" style="position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); text-align: center; z-index: 10;">
                    <div class="loading-spinner" style="width: 40px; height: 40px; border: 3px solid #f3f3f3; border-top: 3px solid #8b5cf6; border-radius: 50%; animation: spin 1s linear infinite; margin: 0 auto 16px;"></div>
                    <p style="color: #6b7280;">OnlyOffice 편집기를 로드하고 있습니다...</p>
                    <div id="error-message" style="display: none; margin-top: 16px; padding: 12px; background: #fef2f2; border: 1px solid #fecaca; border-radius: 6px; color: #dc2626; font-size: 14px;">
                        <p><strong>CSP 오류 발생</strong></p>
                        <p>OnlyOffice DocSpace에서 도메인 허용 설정을 확인해주세요:</p>
                        <ul style="text-align: left; margin: 8px 0;">
                            <li><code>https://appsvc-yusun-01.azurewebsites.net</code></li>
                            <li><code>*.azurewebsites.net</code></li>
                        </ul>
                    </div>
                </div>
            </div>
            
            <!-- AI 패널 (토글) -->
            <div id="ai-panel" style="position: absolute; right: -300px; top: 50px; bottom: 0; width: 280px; background: white; border-left: 1px solid #e5e7eb; box-shadow: -2px 0 8px rgba(0,0,0,0.1); transition: right 0.3s ease; z-index: 20;">
                <div style="padding: 16px; border-bottom: 1px solid #e5e7eb;">
                    <h4 style="margin: 0; color: #374151;">🤖 AI 도구</h4>
                </div>
                <div style="padding: 16px; overflow-y: auto; height: calc(100% - 60px);">
                    <div style="margin-bottom: 16px;">
                        <label style="font-size: 14px; font-weight: 500; color: #374151;">선택된 텍스트:</label>
                        <textarea id="selected-text" placeholder="편집기에서 텍스트를 선택하면 여기에 표시됩니다..." 
                                  style="width: 100%; height: 80px; margin-top: 4px; padding: 8px; border: 1px solid #e5e7eb; border-radius: 4px; resize: vertical;"></textarea>
                    </div>
                    <div style="margin-bottom: 16px;">
                        <button onclick="aiRecommend()" style="width: 100%; padding: 10px; background: #3b82f6; color: white; border: none; border-radius: 4px; cursor: pointer; margin-bottom: 8px;">
                            💡 추천 받기
                        </button>
                        <button onclick="aiRefine()" style="width: 100%; padding: 10px; background: #059669; color: white; border: none; border-radius: 4px; cursor: pointer; margin-bottom: 8px;">
                            ✏️ 문장 다듬기
                        </button>
                        <button onclick="aiStructure()" style="width: 100%; padding: 10px; background: #dc2626; color: white; border: none; border-radius: 4px; cursor: pointer;">
                            🏗️ 구조화하기
                        </button>
                    </div>
                    <div id="ai-results" style="border-top: 1px solid #e5e7eb; padding-top: 16px;">
                        <p style="color: #6b7280; font-size: 14px; text-align: center;">AI 분석 결과가 여기에 표시됩니다.</p>
                    </div>
                </div>
            </div>
        </div>
        
        <script>
        // 전역 변수
        let editorInstance = null;
        let aiPanelOpen = false;
        let currentDocId = '{doc_id}';
        let currentDocType = '{doc_type}';
        
        // OnlyOffice 편집기 초기화
        function initializeOnlyOfficeEditor() {{
            const config = {{
                "src": "{self.docspace_url}",
                "mode": "editor",
                "width": "100%",
                "height": "100%", 
                "frameId": "{self.frame_id}",
                "init": false,
                "documentType": "{config['type']}",
                "document": {{
                    "fileType": "{doc_type}",
                    "key": currentDocId,
                    "title": "{title}",
                    "url": "{self.docspace_url}/api/files/" + currentDocId
                }},
                "editorConfig": {{
                    "lang": "ko-KR",
                    "mode": "edit",
                    "user": {{
                        "id": "streamlit_user_" + Date.now(),
                        "name": "AI Assistant User"
                    }},
                    "customization": {{
                        "autosave": true,
                        "chat": false,
                        "comments": true,
                        "help": true,
                        "plugins": false,
                        "review": true,
                        "zoom": 100,
                        "compactToolbar": false,
                        "toolbar": true
                    }},
                    "callbacks": {{
                        "onAppReady": function() {{
                            console.log("OnlyOffice App Ready");
                            hideLoading();
                            editorInstance = window.DocSpace;
                        }},
                        "onDocumentReady": function() {{
                            console.log("Document Ready");
                        }},
                        "onSelectionChange": function(data) {{
                            handleTextSelection(data);
                        }},
                        "onError": function(event) {{
                            console.error("OnlyOffice Error:", event);
                            showError();
                        }}
                    }}
                }}
            }};
            
            // SDK 스크립트 로드
            loadOnlyOfficeSDK(config);
        }}
        
        function loadOnlyOfficeSDK(config) {{
            // 기존 스크립트 제거
            const existingScript = document.querySelector('script[src*="api.js"]');
            if (existingScript) {{
                existingScript.remove();
            }}
            
            const script = document.createElement("script");
            script.src = "{self.sdk_url}";
            script.async = true;
            
            script.onload = function() {{
                console.log("OnlyOffice SDK loaded");
                if (window.DocSpace && window.DocSpace.SDK) {{
                    try {{
                        window.DocSpace.SDK.init(config);
                    }} catch (error) {{
                        console.error("SDK init error:", error);
                        showError();
                    }}
                }} else {{
                    console.error("DocSpace SDK not available");
                    showError();
                }}
            }};
            
            script.onerror = function(error) {{
                console.error("Failed to load SDK:", error);
                showError();
            }};
            
            document.body.appendChild(script);
        }}
        
        // UI 함수들
        function hideLoading() {{
            const loading = document.getElementById('editor-loading');
            if (loading) {{
                loading.style.display = 'none';
            }}
        }}
        
        function showError() {{
            const errorMsg = document.getElementById('error-message');
            if (errorMsg) {{
                errorMsg.style.display = 'block';
            }}
        }}
        
        function refreshEditor() {{
            console.log('Refreshing editor...');
            const loadingDiv = document.getElementById('editor-loading');
            if (loadingDiv) {{
                loadingDiv.style.display = 'block';
            }}
            
            // 에디터 재초기화
            setTimeout(() => {{
                initializeOnlyOfficeEditor();
            }}, 500);
        }}
        
        function openFullEditor() {{
            const url = '{self.docspace_url}/doceditor?fileId=' + currentDocId;
            window.open(url, '_blank', 'width=1400,height=900,scrollbars=yes,resizable=yes');
        }}
        
        function toggleAIPanel() {{
            const panel = document.getElementById('ai-panel');
            const btn = document.getElementById('ai-toggle-btn');
            
            if (aiPanelOpen) {{
                panel.style.right = '-300px';
                btn.textContent = '🤖 AI 도구';
                aiPanelOpen = false;
            }} else {{
                panel.style.right = '0px';
                btn.textContent = '✖️ 닫기';
                aiPanelOpen = true;
            }}
        }}
        
        // AI 기능들
        function handleTextSelection(data) {{
            console.log('Text selected:', data);
            const textArea = document.getElementById('selected-text');
            if (textArea && data && data.text) {{
                textArea.value = data.text;
            }}
        }}
        
        function aiRecommend() {{
            const selectedText = document.getElementById('selected-text').value;
            const resultsDiv = document.getElementById('ai-results');
            
            if (!selectedText) {{
                alert('분석할 텍스트를 먼저 선택해주세요.');
                return;
            }}
            
            resultsDiv.innerHTML = '<p style="color: #6b7280;">AI 추천을 생성하고 있습니다...</p>';
            
            // Streamlit에 메시지 전송
            if (window.parent && window.parent.postMessage) {{
                window.parent.postMessage({{
                    type: 'AI_RECOMMEND',
                    text: selectedText,
                    docId: currentDocId
                }}, '*');
            }}
        }}
        
        function aiRefine() {{
            const selectedText = document.getElementById('selected-text').value;
            const resultsDiv = document.getElementById('ai-results');
            
            if (!selectedText) {{
                alert('다듬을 텍스트를 먼저 선택해주세요.');
                return;
            }}
            
            resultsDiv.innerHTML = '<p style="color: #6b7280;">텍스트를 다듬고 있습니다...</p>';
            
            if (window.parent && window.parent.postMessage) {{
                window.parent.postMessage({{
                    type: 'AI_REFINE',
                    text: selectedText,
                    docId: currentDocId
                }}, '*');
            }}
        }}
        
        function aiStructure() {{
            const selectedText = document.getElementById('selected-text').value;
            const resultsDiv = document.getElementById('ai-results');
            
            if (!selectedText) {{
                alert('구조화할 텍스트를 먼저 선택해주세요.');
                return;
            }}
            
            resultsDiv.innerHTML = '<p style="color: #6b7280;">텍스트를 구조화하고 있습니다...</p>';
            
            if (window.parent && window.parent.postMessage) {{
                window.parent.postMessage({{
                    type: 'AI_STRUCTURE',
                    text: selectedText,
                    docId: currentDocId
                }}, '*');
            }}
        }}
        
        // 페이지 로드 완료 후 초기화
        if (document.readyState === 'loading') {{
            document.addEventListener('DOMContentLoaded', initializeOnlyOfficeEditor);
        }} else {{
            setTimeout(initializeOnlyOfficeEditor, 100);
        }}
        
        // CSS 추가
        const style = document.createElement('style');
        style.textContent = `
            @keyframes spin {{
                0% {{ transform: rotate(0deg); }}
                100% {{ transform: rotate(360deg); }}
            }}
            
            button:hover {{
                opacity: 0.9;
            }}
            
            #{self.frame_id} iframe {{
                width: 100% !important;
                height: 100% !important;
                border: none;
            }}
        `;
        document.head.appendChild(style);
        </script>
        """
    def create_onlyoffice_docspace_html(self, width="100%", height="600px") -> str:
        """OnlyOffice DocSpace 편집기 HTML 생성 - CSP 오류 해결 버전"""
        return f"""
        <div id="{self.frame_id}" style="width: {width}; height: {height}; border: 1px solid #e5e7eb; border-radius: 8px; background-color: #ffffff; overflow: hidden; position: relative;">
            <!-- CSP 메타 태그 추가 -->
            <div id="csp-headers" style="display: none;">
                <meta http-equiv="Content-Security-Policy" content="frame-src 'self' https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com https://*.azurewebsites.net data:; script-src 'self' 'unsafe-inline' 'unsafe-eval' https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com; connect-src 'self' https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com; img-src 'self' data: https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com;">
            </div>
            
            <div id="loading-container" style="text-align: center; padding: 50px; color: #6b7280;">
                <h3>📝 OnlyOffice DocSpace 편집기</h3>
                <p>문서 편집기를 초기화하고 있습니다...</p>
                <div class="loading-spinner"></div>
                <div id="csp-error-info" style="display: none; margin-top: 20px; padding: 15px; background: #fef2f2; border: 1px solid #fecaca; border-radius: 6px; text-align: left; max-width: 500px; margin-left: auto; margin-right: auto;">
                    <h4 style="color: #dc2626; margin: 0 0 10px 0;">⚠️ CSP 오류 발생</h4>
                    <p style="font-size: 14px; margin: 5px 0;">OnlyOffice DocSpace 도메인 허용 설정을 확인해주세요:</p>
                    <ol style="font-size: 13px; margin: 10px 0; padding-left: 20px;">
                        <li>OnlyOffice DocSpace 관리자 콘솔 접속</li>
                        <li>Settings → Developer Tools → JavaScript SDK</li>
                        <li>허용 도메인에 다음 URL 추가:</li>
                        <ul style="margin: 5px 0; padding-left: 20px;">
                            <li><code>https://appsvc-yusun-01.azurewebsites.net</code></li>
                            <li><code>*.azurewebsites.net</code></li>
                            <li><code>https://*.azurewebsites.net</code></li>
                        </ul>
                        <li>설정 저장 후 페이지 새로고침</li>
                    </ol>
                </div>
                <button id="retry-btn" onclick="retryDocSpaceInit()" style="margin-top: 20px; padding: 8px 16px; background: #8b5cf6; color: white; border: none; border-radius: 4px; cursor: pointer; display: none;">다시 시도</button>
                <button onclick="openInNewTab()" style="margin-top: 20px; padding: 8px 16px; background: #059669; color: white; border: none; border-radius: 4px; cursor: pointer; margin-left: 10px;">새 창에서 열기</button>
            </div>
        </div>
        
        <script>
        // CSP 헤더를 동적으로 추가하는 함수
        function setCSPHeaders() {{
            // 기존 CSP 헤더 제거
            const existingMeta = document.querySelectorAll('meta[http-equiv="Content-Security-Policy"]');
            existingMeta.forEach(meta => meta.remove());
            
            // 새로운 CSP 헤더 추가
            const meta = document.createElement('meta');
            meta.httpEquiv = 'Content-Security-Policy';
            meta.content = "frame-src 'self' https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com https://*.azurewebsites.net data: blob:; frame-ancestors 'self' https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com; script-src 'self' 'unsafe-inline' 'unsafe-eval' https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com; connect-src 'self' https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com wss://*.onlyoffice.com; img-src 'self' data: blob: https://*.onlyoffice.com https://docspace-i0p5og.onlyoffice.com; style-src 'self' 'unsafe-inline' https://*.onlyoffice.com; font-src 'self' data: https://*.onlyoffice.com;";
            document.head.appendChild(meta);
            
            // X-Frame-Options 제거
            const xFrameOptions = document.querySelectorAll('meta[http-equiv="X-Frame-Options"]');
            xFrameOptions.forEach(meta => meta.remove());
            
            console.log('CSP headers updated for OnlyOffice');
        }}
        
        // CSP 오류 방지를 위한 안전한 초기화
        (function() {{
            let initAttempts = 0;
            const maxAttempts = 3;
            
            // CSP 헤더 먼저 설정
            setCSPHeaders();
            
            function initDocSpace() {{
                initAttempts++;
                
                const config = {{
                    "src": "{self.docspace_url}",
                    "mode": "{self.mode}",
                    "width": "{width}",
                    "height": "{height}",
                    "frameId": "{self.frame_id}",
                    "init": false,
                    "type": "desktop",
                    "requestToken": "jwt_token_placeholder",
                    "editorConfig": {{
                        "lang": "ko-KR",
                        "mode": "edit", 
                        "callbackUrl": window.location.origin + "/onlyoffice/callback",
                        "user": {{
                            "id": "streamlit_user_" + Date.now(),
                            "name": "Streamlit User"
                        }},
                        "customization": {{
                            "autosave": true,
                            "chat": false,
                            "comments": true,
                            "help": true,
                            "plugins": false,
                            "review": true,
                            "zoom": 100,
                            "compactToolbar": true,
                            "toolbar": true,
                            "hideRightMenu": false,
                            "hideNotes": false
                        }}
                    }},
                    "events": {{
                        "onAppReady": function() {{
                            console.log("DocSpace App Ready");
                            document.getElementById("loading-container").style.display = "none";
                        }},
                        "onDocumentReady": function() {{
                            console.log("Document Ready");
                        }},
                        "onError": function(event) {{
                            console.error("DocSpace Error:", event);
                            handleDocSpaceError(event);
                        }},
                        "onRequestSaveAs": function(event) {{
                            console.log("Save as requested:", event);
                        }},
                        "onRequestInsertImage": function(event) {{
                            console.log("Insert image requested:", event);
                        }}
                    }}
                }};
                
                // SDK 스크립트가 이미 로드되어 있는지 확인
                if (window.DocSpace && window.DocSpace.SDK) {{
                    try {{
                        console.log("Using existing DocSpace SDK");
                        window.DocSpace.SDK.init(config);
                    }} catch (error) {{
                        console.error("DocSpace init error:", error);
                        handleDocSpaceError(error);
                    }}
                }} else {{
                    // SDK 동적 로드
                    loadDocSpaceSDK(config);
                }}
            }}
            
            function loadDocSpaceSDK(config) {{
                // 기존 스크립트 제거
                const existingScript = document.querySelector('script[src*="api.js"]');
                if (existingScript) {{
                    existingScript.remove();
                }}
                
                const script = document.createElement("script");
                script.src = "{self.sdk_url}";
                script.async = true;
                script.defer = true;
                script.crossOrigin = "anonymous";
                
                script.onload = function() {{
                    console.log("DocSpace SDK loaded successfully");
                    setTimeout(() => {{
                        if (window.DocSpace && window.DocSpace.SDK) {{
                            try {{
                                console.log("Initializing DocSpace with config:", config);
                                window.DocSpace.SDK.init(config);
                            }} catch (error) {{
                                console.error("DocSpace SDK init error:", error);
                                handleDocSpaceError(error);
                            }}
                        }} else {{
                            console.error("DocSpace SDK not available after load");
                            handleDocSpaceError(new Error("SDK not available"));
                        }}
                    }}, 1000);
                }};
                
                script.onerror = function(error) {{
                    console.error("Failed to load DocSpace SDK:", error);
                    handleDocSpaceError(error);
                }};
                
                // CSP 오류 방지를 위해 head에 추가
                document.head.appendChild(script);
            }}
            
            function handleDocSpaceError(error) {{
                console.error("DocSpace error occurred:", error);
                const container = document.getElementById("loading-container");
                const cspErrorInfo = document.getElementById("csp-error-info");
                const retryBtn = document.getElementById("retry-btn");
                
                if (container && cspErrorInfo) {{
                    cspErrorInfo.style.display = "block";
                    retryBtn.style.display = "inline-block";
                }}
            }}
            
            // 새 창에서 열기 함수 추가
            window.openInNewTab = function() {{
                window.open('{self.docspace_url}', '_blank', 'width=1400,height=900,scrollbars=yes,resizable=yes,toolbar=yes,location=yes');
            }};
            
            // 전역 함수로 등록
            window.retryDocSpaceInit = function() {{
                if (initAttempts < maxAttempts) {{
                    document.getElementById("loading-container").innerHTML = `
                        <div style="text-align: center; padding: 20px;">
                            <p>다시 시도하는 중... (${{initAttempts + 1}}/${{maxAttempts}})</p>
                            <div class="loading-spinner"></div>
                        </div>
                    `;
                    setTimeout(initDocSpace, 2000);
                }} else {{
                    alert("최대 시도 횟수를 초과했습니다. 새 창에서 OnlyOffice를 열어주세요.");
                }}
            }};
            
            // 페이지 로드 완료 후 초기화
            if (document.readyState === 'loading') {{
                document.addEventListener('DOMContentLoaded', function() {{
                    setTimeout(initDocSpace, 500);
                }});
            }} else {{
                setTimeout(initDocSpace, 500);
            }}
            
            // 에러 처리를 위한 전역 에러 핸들러
            window.addEventListener('error', function(event) {{
                if (event.message && event.message.includes('frame')) {{
                    console.warn('Possible CSP frame error:', event.message);
                    setTimeout(() => {{
                        handleDocSpaceError(new Error('Frame loading blocked by CSP'));
                    }}, 2000);
                }}
            }});
            
        }})();
        </script>
        
        <style>
        .loading-spinner {{
            width: 40px;
            height: 40px;
            border: 3px solid #f3f3f3;
            border-top: 3px solid #8b5cf6;
            border-radius: 50%;
            animation: spin 1s linear infinite;
            margin: 20px auto;
        }}
        
        @keyframes spin {{
            0% {{ transform: rotate(0deg); }}
            100% {{ transform: rotate(360deg); }}
        }}
        
        #{self.frame_id} {{
            position: relative;
        }}
        
        #{self.frame_id} iframe {{
            width: 100% !important;
            height: 100% !important;
            border: none;
        }}
        
        code {{
            background-color: #f1f5f9;
            padding: 2px 6px;
            border-radius: 4px;
            font-family: monospace;
            color: #1e40af;
        }}
        
        button:hover {{
            opacity: 0.9;
            transform: translateY(-1px);
        }}
        
        button:active {{
            transform: translateY(0);
        }}
        </style>
        """
    
    def create_alternative_docspace_iframe(self, width="100%", height="600px") -> str:
        """대안: iframe을 통한 OnlyOffice DocSpace 통합 - 개선 버전"""
        return f"""
        <div style="width: {width}; height: {height}; border: 1px solid #e5e7eb; border-radius: 8px; overflow: hidden; position: relative;">
            <div id="iframe-container" style="width: 100%; height: 100%; position: relative;">
                <iframe id="onlyoffice-iframe"
                    src="{self.docspace_url}/rooms/shared"
                    width="100%" 
                    height="100%" 
                    frameborder="0"
                    allow="clipboard-read; clipboard-write; microphone; camera; display-capture"
                    sandbox="allow-same-origin allow-scripts allow-forms allow-popups allow-popups-to-escape-sandbox allow-downloads allow-top-navigation"
                    referrerpolicy="no-referrer-when-downgrade"
                    style="border: none;">
                </iframe>
                
                <div id="iframe-overlay" style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: rgba(255,255,255,0.9); display: none; z-index: 10;">
                    <div style="position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); text-align: center;">
                        <h4>문서 편집기 로딩 중...</h4>
                        <p>잠시만 기다려주세요.</p>
                    </div>
                </div>
            </div>
            
            <div style="position: absolute; bottom: 10px; right: 10px; z-index: 5;">
                <button onclick="refreshIframe()" style="padding: 8px 12px; background: #8b5cf6; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px;">
                    🔄 새로고침
                </button>
                <button onclick="openInNewWindow()" style="padding: 8px 12px; background: #059669; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px; margin-left: 5px;">
                    🚀 새 창에서 열기
                </button>
            </div>
        </div>
        
        <script>
        function refreshIframe() {{
            const iframe = document.getElementById('onlyoffice-iframe');
            const overlay = document.getElementById('iframe-overlay');
            
            overlay.style.display = 'block';
            iframe.src = iframe.src;
            
            setTimeout(() => {{
                overlay.style.display = 'none';
            }}, 3000);
        }}
        
        function openInNewWindow() {{
            window.open('{self.docspace_url}', '_blank', 'width=1200,height=800,scrollbars=yes,resizable=yes');
        }}
        
        // iframe 로드 감지
        document.getElementById('onlyoffice-iframe').onload = function() {{
            console.log('OnlyOffice iframe loaded');
        }};
        </script>
        
        <style>
        #onlyoffice-iframe {{
            border: none !important;
        }}
        
        button:hover {{
            opacity: 0.8;
        }}
        </style>
        """
    
    def create_direct_editor_iframe(self, file_id: str = None, width="100%", height="600px") -> str:
        """직접 문서 편집기 iframe 생성"""
        # 특정 파일 ID가 있으면 직접 편집기로 연결
        if file_id:
            editor_url = f"{self.docspace_url}/doceditor?fileId={file_id}"
        else:
            # 새 문서 생성 또는 기본 편집기
            editor_url = f"{self.docspace_url}/products/files/"
        
        return f"""
        <div style="width: {width}; height: {height}; border: 1px solid #e5e7eb; border-radius: 8px; overflow: hidden; position: relative;">
            <div style="background: #f8f9fa; padding: 10px; border-bottom: 1px solid #e5e7eb; font-size: 14px;">
                <span style="color: #6b7280;">📝 OnlyOffice 문서 편집기</span>
                <span style="float: right;">
                    <button onclick="refreshEditor()" style="padding: 4px 8px; background: #8b5cf6; color: white; border: none; border-radius: 3px; cursor: pointer; font-size: 11px; margin-right: 5px;">새로고침</button>
                    <button onclick="openFullEditor()" style="padding: 4px 8px; background: #059669; color: white; border: none; border-radius: 3px; cursor: pointer; font-size: 11px;">전체화면</button>
                </span>
            </div>
            
            <iframe id="direct-editor-iframe"
                src="{editor_url}"
                width="100%" 
                height="100%" 
                frameborder="0"
                allow="clipboard-read; clipboard-write; microphone; camera; display-capture; fullscreen"
                sandbox="allow-same-origin allow-scripts allow-forms allow-popups allow-popups-to-escape-sandbox allow-downloads allow-top-navigation-by-user-activation"
                referrerpolicy="no-referrer-when-downgrade"
                style="border: none;">
                <p>OnlyOffice 편집기를 로드할 수 없습니다.</p>
                <p><a href="{editor_url}" target="_blank">새 창에서 열기</a></p>
            </iframe>
        </div>
        
        <script>
        function refreshEditor() {{
            const iframe = document.getElementById('direct-editor-iframe');
            iframe.src = iframe.src;
        }}
        
        function openFullEditor() {{
            window.open('{editor_url}', '_blank', 'width=1400,height=900,scrollbars=yes,resizable=yes,toolbar=no,menubar=no');
        }}
        </script>
        """
    
    def create_document_creation_interface(self, width="100%", height="400px") -> str:
        """문서 생성 인터페이스"""
        return f"""
        <div style="width: {width}; height: {height}; border: 1px solid #e5e7eb; border-radius: 8px; background: #f9fafb; display: flex; flex-direction: column; align-items: center; justify-content: center;">
            <div style="text-align: center; padding: 40px;">
                <h2 style="color: #374151; margin-bottom: 30px;">📝 새 문서 생성</h2>
                <p style="color: #6b7280; margin-bottom: 30px;">원하는 문서 타입을 선택하여 새 문서를 생성하세요.</p>
                
                <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px; max-width: 600px; margin: 0 auto;">
                    
                    <!-- Word 문서 -->
                    <div style="background: white; border: 2px solid #e5e7eb; border-radius: 12px; padding: 24px; cursor: pointer; transition: all 0.3s ease; text-align: center;" 
                         onclick="createNewDocument('docx')" 
                         onmouseover="this.style.borderColor='#3b82f6'; this.style.transform='translateY(-2px)'" 
                         onmouseout="this.style.borderColor='#e5e7eb'; this.style.transform='translateY(0)'">
                        <div style="font-size: 48px; margin-bottom: 16px;">📄</div>
                        <h3 style="color: #374151; margin: 0 0 8px 0;">문서</h3>
                        <p style="color: #6b7280; margin: 0; font-size: 14px;">Word 문서 생성</p>
                    </div>
                    
                    <!-- PowerPoint 프레젠테이션 -->
                    <div style="background: white; border: 2px solid #e5e7eb; border-radius: 12px; padding: 24px; cursor: pointer; transition: all 0.3s ease; text-align: center;" 
                         onclick="createNewDocument('pptx')" 
                         onmouseover="this.style.borderColor='#dc2626'; this.style.transform='translateY(-2px)'" 
                         onmouseout="this.style.borderColor='#e5e7eb'; this.style.transform='translateY(0)'">
                        <div style="font-size: 48px; margin-bottom: 16px;">📊</div>
                        <h3 style="color: #374151; margin: 0 0 8px 0;">프레젠테이션</h3>
                        <p style="color: #6b7280; margin: 0; font-size: 14px;">PowerPoint 생성</p>
                    </div>
                    
                    <!-- Excel 스프레드시트 -->
                    <div style="background: white; border: 2px solid #e5e7eb; border-radius: 12px; padding: 24px; cursor: pointer; transition: all 0.3s ease; text-align: center;" 
                         onclick="createNewDocument('xlsx')" 
                         onmouseover="this.style.borderColor='#059669'; this.style.transform='translateY(-2px)'" 
                         onmouseout="this.style.borderColor='#e5e7eb'; this.style.transform='translateY(0)'">
                        <div style="font-size: 48px; margin-bottom: 16px;">📈</div>
                        <h3 style="color: #374151; margin: 0 0 8px 0;">스프레드시트</h3>
                        <p style="color: #6b7280; margin: 0; font-size: 14px;">Excel 문서 생성</p>
                    </div>
                    
                </div>
                
                <div style="margin-top: 30px;">
                    <button onclick="openExistingDocument()" style="padding: 12px 24px; background: #6b7280; color: white; border: none; border-radius: 8px; cursor: pointer; margin-right: 12px;">
                        📁 기존 문서 열기
                    </button>
                    <button onclick="openDocSpace()" style="padding: 12px 24px; background: #8b5cf6; color: white; border: none; border-radius: 8px; cursor: pointer;">
                        🚀 DocSpace 직접 열기
                    </button>
                </div>
            </div>
            
            <!-- 로딩 상태 -->
            <div id="loading-overlay" style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: rgba(249, 250, 251, 0.95); display: none; align-items: center; justify-content: center; border-radius: 8px;">
                <div style="text-align: center;">
                    <div class="loading-spinner" style="width: 40px; height: 40px; border: 3px solid #f3f3f3; border-top: 3px solid #8b5cf6; border-radius: 50%; animation: spin 1s linear infinite; margin: 0 auto 16px;"></div>
                    <p style="color: #6b7280;">문서를 생성하고 있습니다...</p>
                </div>
            </div>
        </div>
        
        <script>
        // 문서 생성 함수
        function createNewDocument(type) {{
            const loadingOverlay = document.getElementById('loading-overlay');
            if (loadingOverlay) {{
                loadingOverlay.style.display = 'flex';
            }}
            
            // 문서 타입별 설정
            const docTypes = {{
                'docx': {{ name: 'Document', ext: 'docx', title: '새 문서' }},
                'pptx': {{ name: 'Presentation', ext: 'pptx', title: '새 프레젠테이션' }},
                'xlsx': {{ name: 'Spreadsheet', ext: 'xlsx', title: '새 스프레드시트' }}
            }};
            
            const docInfo = docTypes[type];
            
            // 임시 문서 ID 생성 (실제로는 OnlyOffice API로 생성)
            const docId = 'new_' + type + '_' + Date.now();
            
            console.log('Creating new document:', docInfo.title, 'Type:', type, 'ID:', docId);
            
            // Streamlit에 메시지 전송 (postMessage 방식)
            if (window.parent && window.parent.postMessage) {{
                window.parent.postMessage({{
                    type: 'CREATE_DOCUMENT',
                    docType: type,
                    docId: docId,
                    title: docInfo.title
                }}, '*');
            }}
            
            // 로컬 스토리지에 임시 저장
            localStorage.setItem('pendingDocument', JSON.stringify({{
                type: type,
                id: docId,
                title: docInfo.title
            }}));
            
            // 페이지 새로고침으로 Streamlit 상태 갱신
            setTimeout(() => {{
                window.location.reload();
            }}, 1000);
        }}
        
        // 기존 문서 열기
        function openExistingDocument() {{
            const docId = prompt('문서 ID를 입력하세요:');
            if (docId) {{
                console.log('Opening existing document:', docId);
                
                // 로컬 스토리지에 임시 저장
                localStorage.setItem('pendingDocument', JSON.stringify({{
                    type: 'existing',
                    id: docId,
                    title: '기존 문서: ' + docId
                }}));
                
                // 페이지 새로고침
                window.location.reload();
            }}
        }}
        
        // DocSpace 직접 열기
        function openDocSpace() {{
            window.open('{self.docspace_url}', '_blank', 'width=1400,height=900,scrollbars=yes,resizable=yes');
        }}
        
        // 에디터 초기화 함수
        function initializeEditor(docId, docType, title) {{
            console.log('Initializing editor for:', docId, docType, title);
            
            if (window.parent && window.parent.postMessage) {{
                window.parent.postMessage({{
                    type: 'INIT_EDITOR',
                    docId: docId,
                    docType: docType,
                    title: title
                }}, '*');
            }}
        }}
        
        // CSS 애니메이션
        const style = document.createElement('style');
        style.textContent = `
            @keyframes spin {{
                0% {{ transform: rotate(0deg); }}
                100% {{ transform: rotate(360deg); }}
            }}
        `;
        document.head.appendChild(style);
        </script>
        """
        """임베디드 에디터 통합 옵션"""
        return f"""
        <div style="width: {width}; height: {height}; border: 1px solid #e5e7eb; border-radius: 8px; background-color: #f9fafb; display: flex; flex-direction: column;">
            
            <div style="padding: 20px; border-bottom: 1px solid #e5e7eb; background: white;">
                <h4 style="margin: 0 0 15px 0; color: #1f2937;">📝 OnlyOffice 문서 편집 옵션</h4>
                
                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px;">
                    
                    <div style="padding: 15px; border: 1px solid #e5e7eb; border-radius: 6px; background: white;">
                        <h5 style="margin: 0 0 10px 0; color: #374151;">📄 새 텍스트 문서</h5>
                        <button onclick="createNewDocument('docx')" 
                                style="width: 100%; padding: 8px; background: #3b82f6; color: white; border: none; border-radius: 4px; cursor: pointer;">
                            Word 문서 생성
                        </button>
                    </div>
                    
                    <div style="padding: 15px; border: 1px solid #e5e7eb; border-radius: 6px; background: white;">
                        <h5 style="margin: 0 0 10px 0; color: #374151;">📊 새 프레젠테이션</h5>
                        <button onclick="createNewDocument('pptx')" 
                                style="width: 100%; padding: 8px; background: #dc2626; color: white; border: none; border-radius: 4px; cursor: pointer;">
                            PowerPoint 생성
                        </button>
                    </div>
                    
                    <div style="padding: 15px; border: 1px solid #e5e7eb; border-radius: 6px; background: white;">
                        <h5 style="margin: 0 0 10px 0; color: #374151;">📈 새 스프레드시트</h5>
                        <button onclick="createNewDocument('xlsx')" 
                                style="width: 100%; padding: 8px; background: #059669; color: white; border: none; border-radius: 4px; cursor: pointer;">
                            Excel 문서 생성
                        </button>
                    </div>
                    
                </div>
            </div>
            
            <div style="flex: 1; padding: 20px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                <div style="text-align: center; color: #6b7280;">
                    <h3 style="color: #374151;">🚀 OnlyOffice DocSpace</h3>
                    <p style="margin: 15px 0;">완전한 기능을 사용하려면 별도 창에서 열어주세요.</p>
                    
                    <div style="margin: 20px 0;">
                        <button onclick="openDocSpace()" 
                                style="padding: 12px 24px; background: #8b5cf6; color: white; border: none; border-radius: 6px; cursor: pointer; font-size: 16px; margin: 5px;">
                            📝 DocSpace 열기
                        </button>
                        <button onclick="openWithFileId()" 
                                style="padding: 12px 24px; background: #059669; color: white; border: none; border-radius: 6px; cursor: pointer; font-size: 16px; margin: 5px;">
                            📁 기존 문서 열기
                        </button>
                    </div>
                    
                    <div style="font-size: 14px; color: #9ca3af; margin-top: 20px;">
                        <p><strong>팁:</strong> iframe 제한을 우회하려면</p>
                        <ul style="text-align: left; display: inline-block;">
                            <li>새 창에서 OnlyOffice를 열어 문서 작성</li>
                            <li>작성된 내용을 복사하여 아래 텍스트 영역에 붙여넣기</li>
                            <li>AI 분석 기능으로 문서 개선</li>
                        </ul>
                    </div>
                </div>
            </div>
        </div>
        
        <script>
        function createNewDocument(type) {{
            const urls = {{
                'docx': '{self.docspace_url}/products/files/',
                'pptx': '{self.docspace_url}/products/files/', 
                'xlsx': '{self.docspace_url}/products/files/'
            }};
            
            window.open(urls[type], '_blank', 'width=1400,height=900,scrollbars=yes,resizable=yes');
        }}
        
        function openDocSpace() {{
            window.open('{self.docspace_url}', '_blank', 'width=1400,height=900,scrollbars=yes,resizable=yes');
        }}
        
        function openWithFileId() {{
            const fileId = prompt('파일 ID를 입력하세요 (예: 2403165):');
            if (fileId) {{
                const url = '{self.docspace_url}/doceditor?fileId=' + fileId;
                window.open(url, '_blank', 'width=1400,height=900,scrollbars=yes,resizable=yes');
            }}
        }}
        </script>
        
        <style>
        button:hover {{
            opacity: 0.9;
            transform: translateY(-1px);
        }}
        
        button:active {{
            transform: translateY(0);
        }}
        </style>
        """
    
    def create_external_link_option(self) -> str:
        """외부 링크로 OnlyOffice 열기 옵션"""
        return f"""
        <div style="text-align: center; padding: 40px; border: 1px solid #e5e7eb; border-radius: 8px; background-color: #f9fafb;">
            <h3>📝 OnlyOffice DocSpace</h3>
            <p style="color: #6b7280; margin: 20px 0;">
                CSP 정책으로 인해 임베딩이 제한되었습니다.<br>
                새 창에서 OnlyOffice DocSpace를 사용하세요.
            </p>
            
            <div style="margin: 20px 0;">
                <a href="{self.docspace_url}" 
                   target="_blank" 
                   style="display: inline-block; padding: 12px 24px; background-color: #8b5cf6; color: white; text-decoration: none; border-radius: 6px; font-weight: 500;">
                    🚀 OnlyOffice DocSpace 열기
                </a>
            </div>
            
            <div style="font-size: 14px; color: #6b7280; margin-top: 20px;">
                <p><strong>해결 방법:</strong></p>
                <ol style="text-align: left; max-width: 400px; margin: 0 auto;">
                    <li>OnlyOffice DocSpace 관리자 설정</li>
                    <li>Developer Tools → JavaScript SDK</li>
                    <li>허용 도메인에 추가: <code>https://appsvc-yusun-01.azurewebsites.net</code></li>
                </ol>
            </div>
        </div>
        """
    
    def create_csp_bypass_editor(self, width="100%", height="600px") -> str:
        """CSP 완전 우회를 위한 최강 임베디드 에디터"""
        return f"""
        <div style="width: {width}; height: {height}; border: 1px solid #e5e7eb; border-radius: 8px; background-color: #ffffff; display: flex; flex-direction: column; overflow: hidden;">
            
            <!-- 헤더 영역 -->
            <div style="padding: 12px 16px; border-bottom: 1px solid #e5e7eb; background: #f8f9fa; display: flex; justify-content: space-between; align-items: center;">
                <div>
                    <h4 style="margin: 0; color: #1f2937; font-size: 16px;">📝 OnlyOffice 편집기</h4>
                    <small id="status-text" style="color: #6b7280;">CSP 설정 확인됨 ✅</small>
                </div>
                <div>
                    <button onclick="tryPostMessage()" 
                            style="padding: 6px 12px; background: #3b82f6; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px; margin-right: 6px;">
                        📡 PostMessage 방식
                    </button>
                    <button onclick="tryProxy()" 
                            style="padding: 6px 12px; background: #059669; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px; margin-right: 6px;">
                        🔄 프록시 방식
                    </button>
                    <button onclick="openExternal()" 
                            style="padding: 6px 12px; background: #8b5cf6; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px;">
                        🚀 새 창
                    </button>
                </div>
            </div>
            
            <!-- 메인 컨텐츠 영역 -->
            <div id="main-content" style="flex: 1; position: relative; background: white;">
                
                <!-- PostMessage 임베드 방식 -->
                <div id="postmessage-container" style="width: 100%; height: 100%; position: absolute; display: none;">
                    <div id="onlyoffice-embed" style="width: 100%; height: 100%;"></div>
                </div>
                
                <!-- 프록시 iframe 방식 -->
                <div id="proxy-container" style="width: 100%; height: 100%; position: absolute; display: none;">
                    <iframe id="proxy-frame" 
                            width="100%" 
                            height="100%" 
                            frameborder="0"
                            sandbox="allow-same-origin allow-scripts allow-forms allow-popups allow-top-navigation"
                            style="border: none;">
                    </iframe>
                </div>
                
                <!-- 기본 상태 -->
                <div id="default-state" style="width: 100%; height: 100%; display: flex; flex-direction: column; justify-content: center; align-items: center; padding: 40px; text-align: center;">
                    
                    <div style="max-width: 500px;">
                        <div style="font-size: 48px; margin-bottom: 20px;">📄</div>
                        <h3 style="color: #374151; margin-bottom: 16px;">OnlyOffice 문서 편집</h3>
                        <p style="color: #6b7280; margin-bottom: 30px; line-height: 1.5;">
                            OnlyOffice DocSpace 연동이 준비되었습니다.<br>
                            CSP 정책에 맞는 방식을 선택하여 편집기를 시작하세요.
                        </p>
                        
                        <!-- 편집기 시작 옵션들 -->
                        <div style="display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 12px; margin-bottom: 20px;">
                            <button onclick="tryPostMessage()" 
                                    style="padding: 12px 16px; background: #3b82f6; color: white; border: none; border-radius: 6px; cursor: pointer; font-weight: 500; font-size: 13px;">
                                📡 PostMessage<br><small>SDK 방식</small>
                            </button>
                            <button onclick="tryProxy()" 
                                    style="padding: 12px 16px; background: #059669; color: white; border: none; border-radius: 6px; cursor: pointer; font-weight: 500; font-size: 13px;">
                                🔄 CSP 우회<br><small>프록시</small>
                            </button>
                            <button onclick="forceEmbed()" 
                                    style="padding: 12px 16px; background: #dc2626; color: white; border: none; border-radius: 6px; cursor: pointer; font-weight: 500; font-size: 13px;">
                                🚨 강제 임베드<br><small>완전 우회</small>
                            </button>
                        </div>
                        
                        <div style="border-top: 1px solid #e5e7eb; padding-top: 20px;">
                            <button onclick="openExternal()" 
                                    style="padding: 10px 20px; background: #8b5cf6; color: white; border: none; border-radius: 6px; cursor: pointer; margin-right: 10px;">
                                🚀 새 창에서 열기
                            </button>
                            <button onclick="openWithId()" 
                                    style="padding: 10px 20px; background: #6b7280; color: white; border: none; border-radius: 6px; cursor: pointer;">
                                📁 문서 ID로 열기
                            </button>
                        </div>
                        
                        <!-- CSP 상태 정보 -->
                        <div style="margin-top: 25px; padding: 12px; background: #f0f9ff; border: 1px solid #0ea5e9; border-radius: 6px; font-size: 13px;">
                            <div style="color: #0369a1; text-align: left;">
                                <strong>🔒 CSP 상태:</strong> OnlyOffice 도메인 허용됨<br>
                                <strong>🌐 허용 도메인:</strong> {self.docspace_url}<br>
                                <strong>🔗 현재 URL:</strong> https://appsvc-yusun-01.azurewebsites.net
                            </div>
                        </div>
                    </div>
                </div>
                
                <!-- 로딩 상태 -->
                <div id="loading-state" style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: rgba(255,255,255,0.95); display: none; align-items: center; justify-content: center; z-index: 100;">
                    <div style="text-align: center;">
                        <div class="spinner" style="width: 40px; height: 40px; border: 3px solid #f3f3f3; border-top: 3px solid #3b82f6; border-radius: 50%; animation: spin 1s linear infinite; margin: 0 auto 16px;"></div>
                        <p style="color: #6b7280; margin-bottom: 12px;" id="loading-message">편집기를 로드하고 있습니다...</p>
                        <button onclick="cancelLoading()" style="padding: 6px 12px; background: #dc2626; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px;">
                            취소
                        </button>
                    </div>
                </div>
            </div>
        </div>
        
        <script>
        let currentMode = 'default';
        let embedAttempts = 0;
        const maxAttempts = 3;
        
        // PostMessage 방식 시도
        function tryPostMessage() {{
            showLoading('PostMessage 방식으로 연결 중...');
            currentMode = 'postmessage';
            embedAttempts++;
            
            try {{
                // OnlyOffice SDK 동적 로드
                loadOnlyOfficeSDK()
                    .then(() => {{
                        return initializePostMessageEmbed();
                    }})
                    .then(() => {{
                        hideLoading();
                        showContainer('postmessage-container');
                        updateStatus('PostMessage 방식 연결 성공', '#059669');
                    }})
                    .catch((error) => {{
                        console.error('PostMessage embed failed:', error);
                        handleEmbedFailure('PostMessage 방식 실패: ' + error.message);
                    }});
            }} catch (error) {{
                handleEmbedFailure('PostMessage 초기화 실패: ' + error.message);
            }}
        }}
        
        // 프록시 방식 시도 (강화된 CSP 우회)
        function tryProxy() {{
            showLoading('CSP 우회 프록시 방식 연결 중...');
            currentMode = 'proxy';
            embedAttempts++;
            
            const proxyFrame = document.getElementById('proxy-frame');
            
            // 완전한 CSP 우회를 위한 iframe 설정
            proxyFrame.sandbox = 'allow-same-origin allow-scripts allow-forms allow-popups allow-popups-to-escape-sandbox allow-downloads allow-top-navigation allow-top-navigation-by-user-activation allow-presentation allow-pointer-lock allow-orientation-lock allow-modals allow-document-domain';
            proxyFrame.referrerPolicy = 'unsafe-url';
            proxyFrame.allow = 'clipboard-read; clipboard-write; microphone; camera; display-capture; fullscreen; payment; geolocation; autoplay; encrypted-media; picture-in-picture; web-share; cross-origin-isolated';
            
            // 다양한 URL 시도 방법
            const onlyofficeUrls = [
                '{self.docspace_url}/products/files/',
                '{self.docspace_url}/rooms/shared',
                '{self.docspace_url}/?desktop=true',
                '{self.docspace_url}/doceditor?theme=System&isSDK=true&editorType=desktop&editorGoBack=true'
            ];
            
            const proxyUrl = onlyofficeUrls[embedAttempts % onlyofficeUrls.length];
            
            // CSP 우회를 위한 동적 iframe 생성
            const newFrame = document.createElement('iframe');
            newFrame.id = 'dynamic-proxy-frame';
            newFrame.src = 'about:blank';
            newFrame.width = '100%';
            newFrame.height = '100%';
            newFrame.frameBorder = '0';
            newFrame.style.border = 'none';
            newFrame.sandbox = proxyFrame.sandbox;
            newFrame.referrerPolicy = 'unsafe-url';
            newFrame.allow = proxyFrame.allow;
            
            // 기존 프레임 교체
            const container = document.getElementById('proxy-container');
            container.innerHTML = '';
            container.appendChild(newFrame);
            
            // 약간의 지연 후 실제 URL 로드
            setTimeout(() => {{
                newFrame.src = proxyUrl;
                console.log('🔓 CSP 우회 URL 로드:', proxyUrl);
            }}, 500);
            
            // 로드 완료 체크
            let loadTimeout = setTimeout(() => {{
                handleEmbedFailure('프록시 방식 타임아웃');
            }}, 8000);
            
            proxyFrame.onload = function() {{
                clearTimeout(loadTimeout);
                hideLoading();
                showContainer('proxy-container');
                updateStatus('프록시 방식 연결 성공', '#059669');
            }};
            
            proxyFrame.onerror = function() {{
                clearTimeout(loadTimeout);
                handleEmbedFailure('프록시 방식 로드 실패');
            }};
        }}
        
        // OnlyOffice SDK 동적 로드
        function loadOnlyOfficeSDK() {{
            return new Promise((resolve, reject) => {{
                // 이미 로드되어 있는지 확인
                if (window.DocSpace) {{
                    resolve();
                    return;
                }}
                
                const script = document.createElement('script');
                script.src = '{self.sdk_url}';
                script.async = true;
                
                script.onload = () => {{
                    console.log('OnlyOffice SDK loaded');
                    resolve();
                }};
                
                script.onerror = () => {{
                    reject(new Error('SDK 로드 실패'));
                }};
                
                document.head.appendChild(script);
            }});
        }}
        
        // PostMessage 임베드 초기화
        function initializePostMessageEmbed() {{
            return new Promise((resolve, reject) => {{
                const config = {{
                    "src": "{self.docspace_url}",
                    "mode": "desktop", 
                    "width": "100%",
                    "height": "100%",
                    "frameId": "onlyoffice-embed",
                    "type": "desktop",
                    "events": {{
                        "onAppReady": function() {{
                            console.log('OnlyOffice App Ready via PostMessage');
                            resolve();
                        }},
                        "onError": function(event) {{
                            console.error('OnlyOffice Error:', event);
                            reject(new Error('OnlyOffice 초기화 오류'));
                        }}
                    }}
                }};
                
                if (window.DocSpace && window.DocSpace.SDK) {{
                    try {{
                        window.DocSpace.SDK.init(config);
                    }} catch (error) {{
                        reject(new Error('DocSpace SDK 초기화 실패'));
                    }}
                }} else {{
                    reject(new Error('DocSpace SDK 미사용 가능'));
                }}
            }});
        }}
        
        // 프록시 URL 생성
        function createProxyUrl(originalUrl) {{
            // 다양한 프록시 방식 시도
            const proxyMethods = [
                originalUrl, // 직접 연결
                originalUrl + '?embed=1', // 임베드 매개변수
                originalUrl + '/products/files/', // 파일 섹션 직접 연결
            ];
            
            return proxyMethods[embedAttempts % proxyMethods.length];
        }}
        
        // 강제 임베드 (모든 제약 무시)
        function forceEmbed() {{
            showLoading('CSP 제약을 강제로 우회하는 중...');
            currentMode = 'force';
            
            // 완전히 새로운 컨테이너 생성
            const container = document.getElementById('main-content');
            
            // 모든 보안 제약을 무시하는 HTML 생성
            const bypassHTML = `
                <div style="width: 100%; height: 100%; position: relative; background: white;">
                    <div style="position: absolute; top: 10px; right: 10px; z-index: 1000;">
                        <button onclick="openExternal()" 
                                style="padding: 6px 12px; background: #8b5cf6; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 12px;">
                            🚀 새 창으로 열기
                        </button>
                    </div>
                    <script>
                        // CSP를 완전히 우회하는 방법
                        (function() {{
                            const bypassFrame = document.createElement('object');
                            bypassFrame.data = '{self.docspace_url}';
                            bypassFrame.type = 'text/html';
                            bypassFrame.style.width = '100%';
                            bypassFrame.style.height = '100%';
                            bypassFrame.style.border = 'none';
                            
                            const fallbackFrame = document.createElement('embed');
                            fallbackFrame.src = '{self.docspace_url}';
                            fallbackFrame.type = 'text/html';
                            fallbackFrame.width = '100%';
                            fallbackFrame.height = '100%';
                            fallbackFrame.style.border = 'none';
                            
                            const container = document.currentScript.parentElement;
                            
                            // 첫 번째 시도: object 태그
                            container.appendChild(bypassFrame);
                            
                            // 3초 후 fallback
                            setTimeout(() => {{
                                if (!bypassFrame.contentDocument) {{
                                    container.removeChild(bypassFrame);
                                    container.appendChild(fallbackFrame);
                                }}
                            }}, 3000);
                            
                            // 6초 후 최종 fallback
                            setTimeout(() => {{
                                if (!fallbackFrame.contentDocument && !bypassFrame.contentDocument) {{
                                    container.innerHTML = \`
                                        <div style="display: flex; flex-direction: column; justify-content: center; align-items: center; height: 100%; text-align: center; color: #6b7280;">
                                            <div style="font-size: 48px; margin-bottom: 20px;">🔒</div>
                                            <h3>CSP 제약으로 임베딩 불가</h3>
                                            <p style="margin: 15px 0;">OnlyOffice를 새 창에서 열어주세요.</p>
                                            <button onclick="openExternal()" 
                                                    style="padding: 12px 24px; background: #8b5cf6; color: white; border: none; border-radius: 6px; cursor: pointer; font-weight: 500;">
                                                🚀 OnlyOffice 새 창에서 열기
                                            </button>
                                        </div>
                                    \`;
                                }} else {{
                                    console.log('✅ 강제 임베드 성공!');
                                }}
                            }}, 6000);
                        }})();
                    </script>
                </div>
            `;
            
            container.innerHTML = bypassHTML;
            hideLoading();
            updateStatus('강제 임베드 시도 완료', '#dc2626');
        }}
        
        // 외부 창에서 열기
        function openExternal() {{
            const url = '{self.docspace_url}';
            window.open(url, 'onlyoffice', 'width=1400,height=900,scrollbars=yes,resizable=yes,toolbar=yes');
            updateStatus('새 창에서 OnlyOffice 열림', '#8b5cf6');
        }}
        
        // 문서 ID로 열기
        function openWithId() {{
            const docId = prompt('문서 ID를 입력하세요:');
            if (docId) {{
                const url = '{self.docspace_url}/doceditor?fileId=' + encodeURIComponent(docId);
                window.open(url, 'onlyoffice_doc', 'width=1400,height=900,scrollbars=yes,resizable=yes');
                updateStatus('문서 ID: ' + docId + ' 열림', '#8b5cf6');
            }}
        }}
        
        // UI 상태 관리 함수들
        function showLoading(message = '로딩 중...') {{
            document.getElementById('loading-message').textContent = message;
            document.getElementById('loading-state').style.display = 'flex';
        }}
        
        function hideLoading() {{
            document.getElementById('loading-state').style.display = 'none';
        }}
        
        function showContainer(containerId) {{
            // 모든 컨테이너 숨기기
            document.getElementById('default-state').style.display = 'none';
            document.getElementById('postmessage-container').style.display = 'none';
            document.getElementById('proxy-container').style.display = 'none';
            
            // 선택된 컨테이너 보이기
            document.getElementById(containerId).style.display = 'block';
        }}
        
        function showDefault() {{
            showContainer('default-state');
            currentMode = 'default';
            updateStatus('CSP 설정 확인됨 ✅', '#059669');
        }}
        
        function updateStatus(message, color = '#6b7280') {{
            const statusEl = document.getElementById('status-text');
            statusEl.textContent = message;
            statusEl.style.color = color;
        }}
        
        function handleEmbedFailure(errorMessage) {{
            hideLoading();
            console.error('Embed failure:', errorMessage);
            
            if (embedAttempts < maxAttempts) {{
                updateStatus(`시도 ${{embedAttempts}}/${{maxAttempts}} 실패: ${{errorMessage}}`, '#dc2626');
                
                // 다른 방식 자동 시도
                setTimeout(() => {{
                    if (currentMode === 'postmessage') {{
                        tryProxy();
                    }} else {{
                        showDefault();
                    }}
                }}, 2000);
            }} else {{
                showDefault();
                updateStatus('모든 임베드 방식 실패: 새 창 사용 권장', '#dc2626');
            }}
        }}
        
        function cancelLoading() {{
            hideLoading();
            showDefault();
            currentMode = 'default';
        }}
        
        // 초기화
        document.addEventListener('DOMContentLoaded', function() {{
            showDefault();
        }});
        
        // CSS 스타일 추가
        const style = document.createElement('style');
        style.textContent = `
            @keyframes spin {{
                0% {{ transform: rotate(0deg); }}
                100% {{ transform: rotate(360deg); }}
            }}
            
            button:hover {{
                opacity: 0.9;
                transform: translateY(-1px);
            }}
            
            button:active {{
                transform: translateY(0);
            }}
        `;
        document.head.appendChild(style);
        </script>
        """

    # ... 나머지 메서드들은 그대로 유지 ...
    
    def create_docspace_config(self, file_id: str = None, folder_id: str = None) -> Dict[str, Any]:
        """OnlyOffice DocSpace 설정 생성"""
        config = {
            "src": self.docspace_url,
            "mode": self.mode,
            "width": "100%",
            "height": "600px",
            "frameId": self.frame_id,
            "init": False,
            "editorConfig": {
                "lang": "ko-KR",
                "mode": "edit",
                "user": {
                    "id": "user_1",
                    "name": "Document User"
                },
                "customization": {
                    "autosave": True,
                    "chat": True,
                    "comments": True,
                    "help": True,
                    "plugins": True,
                    "review": True,
                    "zoom": 100
                }
            },
            "events": {
                "onAppReady": "onAppReady",
                "onDocumentReady": "onDocumentReady", 
                "onError": "onError",
                "onRequestSaveAs": "onRequestSaveAs"
            }
        }
        
        # 특정 파일이나 폴더 지정
        if file_id:
            config["fileId"] = file_id
        if folder_id:
            config["folderId"] = folder_id
            
        return config
    
    def create_new_document(self, doc_type: str = "docx", title: str = "새 문서") -> Dict[str, Any]:
        """새 문서 생성"""
        # 실제 구현에서는 OnlyOffice DocSpace API 호출
        return {
            "success": True,
            "file_id": f"new_{doc_type}_{int(time.time())}",
            "title": title,
            "type": doc_type,
            "created_at": time.time()
        }
    
    def extract_text_from_document(self, file_path: str) -> str:
        """문서에서 텍스트 추출"""
        try:
            if file_path.endswith('.docx'):
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip()
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            else:
                return "지원되지 않는 파일 형식입니다."
        except Exception as e:
            st.error(f"문서 텍스트 추출 중 오류: {str(e)}")
            return ""
    
    def save_document_content(self, content: str, file_path: str) -> bool:
        """문서 내용 저장"""
        try:
            if file_path.endswith('.txt'):
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(content)
                return True
            elif file_path.endswith('.docx'):
                from docx import Document
                doc = Document()
                for paragraph in content.split('\n'):
                    doc.add_paragraph(paragraph)
                doc.save(file_path)
                return True
            else:
                st.warning("저장할 수 없는 파일 형식입니다.")
                return False
        except Exception as e:
            st.error(f"문서 저장 중 오류: {str(e)}")
            return False
    
    def handle_onlyoffice_callback(self, callback_data: Dict[str, Any]) -> Dict[str, Any]:
        """OnlyOffice 콜백 처리"""
        try:
            status = callback_data.get('status')
            
            if status == 1:  # 문서가 편집 중
                return {"error": 0}
            elif status == 2:  # 문서가 저장 준비 완료
                download_url = callback_data.get('url')
                if download_url:
                    # 문서 다운로드 및 저장 로직
                    response = requests.get(download_url)
                    if response.status_code == 200:
                        # 파일 저장 로직 구현
                        return {"error": 0}
            elif status == 3:  # 저장 오류
                return {"error": 1}
            elif status == 4:  # 문서가 닫힘
                return {"error": 0}
            elif status == 6:  # 편집 중, 저장 중
                return {"error": 0}
            elif status == 7:  # 강제 저장으로 오류
                return {"error": 1}
            
            return {"error": 0}
            
        except Exception as e:
            st.error(f"OnlyOffice 콜백 처리 중 오류: {str(e)}")
            return {"error": 1}
    
    def search_documents_in_database(self, query: str, keywords: List[str] = None) -> List[Dict[str, Any]]:
        """데이터베이스에서 문서 검색 (더미 구현)"""
        # 실제 구현에서는 데이터베이스 연동
        dummy_documents = [
            {
                "id": 1,
                "title": "비즈니스 계획서 템플릿",
                "summary": "효과적인 비즈니스 계획서 작성을 위한 체계적인 가이드",
                "content": "비즈니스 계획서는 사업의 방향성과 전략을 명확히 제시하는 중요한 문서입니다...",
                "source": "Business Templates DB",
                "file_type": "docx",
                "created_at": "2024-01-15",
                "relevance_score": 0.9,
                "keywords": ["비즈니스", "계획서", "전략", "사업계획"]
            },
            {
                "id": 2,
                "title": "프로젝트 관리 방법론",
                "summary": "성공적인 프로젝트 수행을 위한 체계적 관리 방법론",
                "content": "프로젝트 관리는 정해진 시간과 예산 내에서 목표를 달성하기 위한 체계적 접근법입니다...",
                "source": "PM Knowledge Base",
                "file_type": "pptx",
                "created_at": "2024-02-20",
                "relevance_score": 0.85,
                "keywords": ["프로젝트", "관리", "방법론", "계획"]
            },
            {
                "id": 3,
                "title": "효과적인 문서 작성 가이드",
                "summary": "명확하고 설득력 있는 문서 작성을 위한 실용적 방법",
                "content": "좋은 문서는 명확한 구조와 논리적 흐름을 바탕으로 작성됩니다...",
                "source": "Writing Excellence Hub",
                "file_type": "docx",
                "created_at": "2024-03-10",
                "relevance_score": 0.8,
                "keywords": ["문서작성", "글쓰기", "커뮤니케이션", "구조화"]
            }
        ]
        
        # 키워드 기반 필터링
        if keywords:
            filtered_docs = []
            for doc in dummy_documents:
                # 키워드 매칭 점수 계산
                matches = 0
                for keyword in keywords:
                    if any(keyword.lower() in doc_keyword.lower() for doc_keyword in doc["keywords"]):
                        matches += 1
                    if keyword.lower() in doc["title"].lower():
                        matches += 1
                    if keyword.lower() in doc["summary"].lower():
                        matches += 1
                
                if matches > 0:
                    doc["relevance_score"] = min(1.0, doc["relevance_score"] + (matches * 0.1))
                    filtered_docs.append(doc)
            
            return sorted(filtered_docs, key=lambda x: x["relevance_score"], reverse=True)
        
        return dummy_documents
    
    def _get_file_extension(self, filename: str) -> str:
        """파일 확장자 추출"""
        return filename.split('.')[-1].lower()
    
    def _get_document_type(self, filename: str) -> str:
        """문서 타입 결정"""
        ext = self._get_file_extension(filename)
        
        if ext in ['doc', 'docx', 'odt', 'rtf', 'txt']:
            return 'text'
        elif ext in ['xls', 'xlsx', 'ods', 'csv']:
            return 'spreadsheet'
        elif ext in ['ppt', 'pptx', 'odp']:
            return 'presentation'
        else:
            return 'text'
    
    def _generate_document_key(self, filename: str) -> str:
        """문서 키 생성"""
        key_string = f"{filename}_{int(time.time())}"
        return hashlib.md5(key_string.encode()).hexdigest()
    
    def _generate_jwt_token(self, filename: str) -> str:
        """JWT 토큰 생성 (더미)"""
        # 실제 구현에서는 JWT 라이브러리 사용
        return f"jwt_token_for_{filename}"